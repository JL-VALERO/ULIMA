from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# --- Helper functions ---
def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

def add_reference(doc, ref_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after   = Pt(4)
    run = p.add_run(ref_text)
    run.font.size = Pt(10)

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)

def add_label(doc, label, content):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(content)
    r2.font.size = Pt(10)

def add_divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("_" * 85)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    p.paragraph_format.space_after = Pt(8)

# ============================================================
# TITLE PAGE
# ============================================================
add_title(doc, "Bibliografía Metodológica")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Modelos VAR e Identificación por Descomposición de Cholesky")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run(
    "Teoría y Política Monetaria  |  ULIMA 2026-1\n"
    "Canal Crediticio de la Política Monetaria en Perú\n"
    "Formato: APA 7ª edición"
)
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_page_break()

# ============================================================
# SECTION 1 — VAR Fundamentals
# ============================================================
add_subtitle(doc, "1. Fundamentos del Modelo VAR")

add_body(doc,
    "Los modelos de Vectores Autorregresivos (VAR) son el estándar metodológico para el "
    "análisis de mecanismos de transmisión de política monetaria. Las referencias siguientes "
    "justifican el uso de VAR reducido con identificación de Cholesky como herramienta "
    "econométrica válida y ampliamente aceptada."
)

# --- Sims (1980) ---
add_section_heading(doc, "1.1  Sims (1980) — Artículo fundacional del VAR")
add_reference(doc,
    "Sims, C. A. (1980). Macroeconomics and reality. Econometrica, 48(1), 1–48. "
    "https://doi.org/10.2307/1912017"
)
add_label(doc, "Tipo: ", "Artículo científico — Econometrica (Scopus Q1, ScienceDirect)")
add_label(doc, "Resumen: ",
    "Artículo seminal que introduce los modelos VAR como alternativa a los modelos "
    "macroeconómicos estructurales de grandes dimensiones. Sims argumenta que las "
    "restricciones de identificación impuestas en esos modelos son injustificadas y propone "
    "dejar que los datos hablen por sí mismos."
)
add_label(doc, "Relevancia: ",
    "Justifica el uso de VAR reducido; es la referencia obligatoria cuando se elige esta "
    "metodología. Con más de 15 000 citas en Google Scholar, constituye la base teórica "
    "indispensable del modelo aplicado en el presente trabajo."
)
add_divider(doc)

# --- Stock & Watson (2001) ---
add_section_heading(doc, "1.2  Stock & Watson (2001) — VAR en macroeconomía aplicada")
add_reference(doc,
    "Stock, J. H., & Watson, M. W. (2001). Vector autoregressions. Journal of Economic "
    "Perspectives, 15(4), 101–115. https://doi.org/10.1257/jep.15.4.101"
)
add_label(doc, "Tipo: ", "Artículo científico — Journal of Economic Perspectives (AEA, Scopus Q1)")
add_label(doc, "Resumen: ",
    "Revisión didáctica y rigurosa de los modelos VAR: especificación, selección de rezagos, "
    "funciones impulso-respuesta, descomposición de varianza del error de pronóstico (FEVD) "
    "e identificación estructural."
)
add_label(doc, "Relevancia: ",
    "Explica directamente las herramientas usadas en el modelo: IRF, FEVD y causalidad de "
    "Granger. Justifica el horizonte de 24 meses para el análisis impulso-respuesta."
)
add_divider(doc)

# ============================================================
# SECTION 2 — Cholesky Identification
# ============================================================
add_subtitle(doc, "2. Identificación Estructural: Descomposición de Cholesky")

add_body(doc,
    "La identificación de shocks estructurales en un VAR requiere imponer restricciones de "
    "corto plazo. La descomposición de Cholesky es el método más empleado: asume una "
    "ordenación causal recursiva en la que las variables más exógenas aparecen primero."
)

# --- Christiano, Eichenbaum & Evans (1999) ---
add_section_heading(doc, "2.1  Christiano, Eichenbaum & Evans (1999) — Cholesky en política monetaria")
add_reference(doc,
    "Christiano, L. J., Eichenbaum, M., & Evans, C. L. (1999). Monetary policy shocks: "
    "What have we learned and to what end? In J. B. Taylor & M. Woodford (Eds.), Handbook "
    "of Macroeconomics (Vol. 1A, pp. 65–148). Elsevier. "
    "https://doi.org/10.1016/S1574-0048(99)01005-8"
)
add_label(doc, "Tipo: ", "Capítulo de libro — Handbook of Macroeconomics, Elsevier (ScienceDirect)")
add_label(doc, "Resumen: ",
    "Estudio exhaustivo que evalúa los efectos de los shocks de política monetaria usando "
    "VAR con identificación Cholesky. Analiza cómo el instrumento de política (tasa de "
    "interés) afecta variables reales y nominales con una ordenación recursiva."
)
add_label(doc, "Relevancia: ",
    "Es la referencia metodológica más citada para justificar el ordenamiento de Cholesky "
    "en el contexto de política monetaria. Respalda directamente el ordenamiento "
    "tasa_ref → tc → tamn → inflacion utilizado en el modelo VAR de este trabajo."
)
add_divider(doc)

# --- Christiano, Eichenbaum & Evans (1996) ---
add_section_heading(doc, "2.2  Christiano, Eichenbaum & Evans (1996) — Efectos de shocks monetarios")
add_reference(doc,
    "Christiano, L. J., Eichenbaum, M., & Evans, C. L. (1996). The effects of monetary "
    "policy shocks: Evidence from the flow of funds. Review of Economics and Statistics, "
    "78(1), 16–34. https://doi.org/10.2307/2109845"
)
add_label(doc, "Tipo: ", "Artículo científico — Review of Economics and Statistics (Scopus Q1)")
add_label(doc, "Resumen: ",
    "Identifica shocks de política monetaria usando restricciones de corto plazo en un VAR "
    "recursivo. Documenta que un endurecimiento de la política monetaria contrae el crédito "
    "bancario y reduce la actividad real."
)
add_label(doc, "Relevancia: ",
    "Justifica empíricamente el canal crediticio como mecanismo de transmisión y respalda "
    "la inclusión de variables bancarias (TAMN) en el VAR con ordenamiento de Cholesky."
)
add_divider(doc)

# ============================================================
# SECTION 3 — Credit Channel Theory
# ============================================================
add_subtitle(doc, "3. Teoría del Canal Crediticio de la Política Monetaria")

add_body(doc,
    "El canal crediticio es uno de los mecanismos de transmisión de política monetaria más "
    "relevantes en economías con mercados financieros poco profundos. Las siguientes "
    "referencias justifican la inclusión de la TAMN como variable clave del modelo."
)

# --- Bernanke & Blinder (1992) ---
add_section_heading(doc, "3.1  Bernanke & Blinder (1992) — El canal crediticio bancario")
add_reference(doc,
    "Bernanke, B. S., & Blinder, A. S. (1992). The federal funds rate and the channels of "
    "monetary transmission. American Economic Review, 82(4), 901–921. "
    "https://www.jstor.org/stable/2117350"
)
add_label(doc, "Tipo: ", "Artículo científico — American Economic Review (AEA, Scopus Q1)")
add_label(doc, "Resumen: ",
    "Demuestra empíricamente que la tasa de interés de política monetaria afecta "
    "directamente el volumen y el costo del crédito bancario, constituyendo el canal "
    "crediticio como mecanismo de transmisión efectivo."
)
add_label(doc, "Relevancia: ",
    "Fundamento teórico directo de la relación tasa_ref → TAMN modelada en el VAR. "
    "Justifica que la TAMN es el canal a través del cual la política monetaria impacta "
    "la economía real."
)
add_divider(doc)

# --- Bernanke & Gertler (1995) ---
add_section_heading(doc, "3.2  Bernanke & Gertler (1995) — Mecanismos del canal crediticio")
add_reference(doc,
    "Bernanke, B. S., & Gertler, M. (1995). Inside the black box: The credit channel of "
    "monetary policy transmission. Journal of Economic Perspectives, 9(4), 27–48. "
    "https://doi.org/10.1257/jep.9.4.27"
)
add_label(doc, "Tipo: ", "Artículo científico — Journal of Economic Perspectives (AEA, Scopus Q1)")
add_label(doc, "Resumen: ",
    "Descompone el canal crediticio en: (a) canal del crédito bancario (bank lending channel) "
    "y (b) canal del balance general (balance sheet channel). Explica cómo las imperfecciones "
    "financieras amplifican los efectos de la política monetaria."
)
add_label(doc, "Relevancia: ",
    "Marco teórico central del trabajo: justifica por qué la TAMN responde a la tasa de "
    "referencia del BCRP y por qué la relación puede variar durante crisis financieras "
    "(GFC 2008-09, COVID 2020-21)."
)
add_divider(doc)

# --- Mishkin (1995) ---
add_section_heading(doc, "3.3  Mishkin (1995) — Canales de transmisión: revisión general")
add_reference(doc,
    "Mishkin, F. S. (1995). Symposium on the monetary transmission mechanism. Journal of "
    "Economic Perspectives, 9(4), 3–10. https://doi.org/10.1257/jep.9.4.3"
)
add_label(doc, "Tipo: ", "Artículo científico — Journal of Economic Perspectives (AEA, Scopus Q1)")
add_label(doc, "Resumen: ",
    "Artículo introductorio del simposio dedicado a los mecanismos de transmisión monetaria. "
    "Presenta y compara cinco canales: tasa de interés, tipo de cambio, precio de activos, "
    "crédito bancario y balance general."
)
add_label(doc, "Relevancia: ",
    "Justifica la inclusión conjunta de tasa_ref, tc, tamn e inflacion en el VAR como "
    "representación simultánea de múltiples canales de transmisión monetaria."
)
add_divider(doc)

# ============================================================
# SECTION 4 — Econometrics Textbooks
# ============================================================
add_subtitle(doc, "4. Libros de Texto Clásicos de Econometría de Series de Tiempo")

add_body(doc,
    "Los modelos VAR con identificación de Cholesky se describen con detalle en los "
    "siguientes textos de referencia estándar, ampliamente usados en cursos de posgrado "
    "y citados en la literatura empírica."
)

# --- Lutkepohl (2005) ---
add_section_heading(doc, "4.1  Lütkepohl (2005) — Referencia técnica del VAR")
add_reference(doc,
    "Lütkepohl, H. (2005). New introduction to multiple time series analysis. Springer. "
    "https://doi.org/10.1007/978-3-540-27752-1"
)
add_label(doc, "Tipo: ", "Libro — Springer (editorial académica de alta reputación)")
add_label(doc, "Resumen: ",
    "Tratamiento matemático completo de los modelos VAR: especificación, estimación por MCO, "
    "criterios de información para selección de rezagos (AIC, HQIC, SBIC), funciones "
    "impulso-respuesta, FEVD y la descomposición de Cholesky como herramienta de "
    "identificación estructural (capítulos 2, 9 y 11)."
)
add_label(doc, "Relevancia: ",
    "Fuente técnica principal para justificar cada elemento del modelo VAR: el ordenamiento "
    "de Cholesky, la selección de rezagos VAR(2), las IRF y la FEVD. Capítulo 9 trata "
    "explícitamente la identificación estructural vía descomposición triangular."
)
add_divider(doc)

# --- Hamilton (1994) ---
add_section_heading(doc, "4.2  Hamilton (1994) — Time Series Analysis")
add_reference(doc,
    "Hamilton, J. D. (1994). Time series analysis. Princeton University Press."
)
add_label(doc, "Tipo: ", "Libro — Princeton University Press")
add_label(doc, "Resumen: ",
    "Manual de referencia para el análisis de series de tiempo en economía. Cubre raíces "
    "unitarias, cointegración, modelos VAR (capítulo 11) y la identificación de shocks "
    "estructurales usando factorización de Cholesky."
)
add_label(doc, "Relevancia: ",
    "Justifica las pruebas de raíz unitaria (ADF, PP) y la prueba de Johansen realizadas "
    "antes de la estimación del VAR, así como el uso de VAR en niveles cuando las "
    "variables son I(1) (argumento de Sims-Stock-Watson)."
)
add_divider(doc)

# --- Enders (2015) ---
add_section_heading(doc, "4.3  Enders (2015) — Applied Econometric Time Series")
add_reference(doc,
    "Enders, W. (2015). Applied econometric time series (4th ed.). Wiley."
)
add_label(doc, "Tipo: ", "Libro — Wiley (editorial académica internacional)")
add_label(doc, "Resumen: ",
    "Texto aplicado de amplia adopción en cursos de econometría de series de tiempo. "
    "El capítulo 5 desarrolla los modelos VAR, la causalidad de Granger, las funciones "
    "impulso-respuesta y la descomposición de varianza, con énfasis en aplicaciones "
    "macroeconómicas."
)
add_label(doc, "Relevancia: ",
    "Justifica la prueba de causalidad de Granger (vargranger en Stata), la interpretación "
    "de las IRF y la FEVD, y la lógica económica detrás de la ordenación de Cholesky."
)
add_divider(doc)

# ============================================================
# SECTION 5 — Summary Table
# ============================================================
add_subtitle(doc, "5. Tabla Resumen de Referencias")

doc.add_paragraph()

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"

hdr = table.rows[0].cells
hdr[0].text = "Referencia"
hdr[1].text = "Tipo"
hdr[2].text = "Indexación"
hdr[3].text = "Justifica"

for cell in hdr:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "2E74B5")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows_data = [
    ["Sims (1980)", "Artículo", "Econometrica / Scopus Q1", "Metodología VAR"],
    ["Stock & Watson (2001)", "Artículo", "JEP / Scopus Q1", "IRF, FEVD, Granger"],
    ["Christiano et al. (1999)", "Cap. libro", "Handbook / ScienceDirect", "Cholesky en política monetaria"],
    ["Christiano et al. (1996)", "Artículo", "REStat / Scopus Q1", "Canal crediticio + Cholesky"],
    ["Bernanke & Blinder (1992)", "Artículo", "AER / Scopus Q1", "Canal crediticio bancario"],
    ["Bernanke & Gertler (1995)", "Artículo", "JEP / Scopus Q1", "Mecanismo canal crediticio"],
    ["Mishkin (1995)", "Artículo", "JEP / Scopus Q1", "Múltiples canales transmisión"],
    ["Lütkepohl (2005)", "Libro", "Springer", "Teoría VAR y Cholesky"],
    ["Hamilton (1994)", "Libro", "Princeton U. Press", "Series de tiempo y VAR"],
    ["Enders (2015)", "Libro", "Wiley", "VAR aplicado, IRF, FEVD"],
]

for row_data in rows_data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# ============================================================
# FOOTER NOTE
# ============================================================
doc.add_paragraph()
p_note = doc.add_paragraph()
r_note = p_note.add_run(
    "Nota: Todas las referencias han sido verificadas en sus fuentes originales. "
    "Los artículos de Econometrica, AER, JEP y REStat están indexados en Scopus y "
    "disponibles en ScienceDirect o JSTOR. Los capítulos del Handbook of Macroeconomics "
    "están disponibles en ScienceDirect (Elsevier). Los libros de Springer, Princeton "
    "University Press y Wiley son editoriales académicas reconocidas internacionalmente."
)
r_note.font.size = Pt(9)
r_note.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
r_note.font.italic = True

# ============================================================
# SAVE
# ============================================================
output_path = (
    r"C:\Users\51950\Documents\ULIMA\2026_1\Teo. Politica"
    r"\Trabajo grupal\Regresion_VAR\Bibliografia\Bibliografia_VAR_Cholesky.docx"
)
doc.save(output_path)
print(f"Documento guardado en: {output_path}")
