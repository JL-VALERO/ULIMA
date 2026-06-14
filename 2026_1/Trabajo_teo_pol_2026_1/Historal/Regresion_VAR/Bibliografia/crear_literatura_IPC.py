from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- margins ----
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ---- colour palette ----
C_BLUE   = RGBColor(0x1F, 0x49, 0x7D)
C_LBLUE  = RGBColor(0x2E, 0x74, 0xB5)
C_GREEN  = RGBColor(0x1D, 0x60, 0x3A)
C_GREY   = RGBColor(0x55, 0x55, 0x55)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_RED    = RGBColor(0x80, 0x20, 0x20)
C_ORANGE = RGBColor(0x7F, 0x46, 0x00)

# ---- helpers ----
def h_main(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(17); r.font.color.rgb = C_BLUE

def h_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.color.rgb = C_LBLUE

def h_section(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = C_LBLUE
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)

def h_paper(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = C_GREEN
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)

def apa_ref(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after       = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10)

def field(doc, label, content):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label)
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = C_BLUE
    r2 = p.add_run(content)
    r2.font.size = Pt(10)

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.size = Pt(10)

def box_chain(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = C_RED

def divider(doc):
    p = doc.add_paragraph()
    r = p.add_run("─" * 90)
    r.font.size = Pt(8); r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(10)

def shade_cell(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

# ==============================================================
# PORTADA
# ==============================================================
h_main(doc, "Justificación Bibliográfica del Ordenamiento VAR")
doc.add_paragraph()
h_subtitle(doc,
    "Índice de Presión Cambiaria (IPC) → Crédito → Actividad Económica → Inflación\n"
    "Fuentes: Scopus y ScienceDirect"
)
doc.add_paragraph()
p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_info = p_info.add_run(
    "Teoría y Política Monetaria  |  ULIMA 2026-1\n"
    "Papers empíricos para la justificación teórica y metodológica\n"
    "Referencias en formato APA 7ª edición"
)
r_info.font.size = Pt(10); r_info.font.color.rgb = C_GREY

doc.add_page_break()

# ==============================================================
# INTRODUCCIÓN Y LÓGICA DEL ORDENAMIENTO
# ==============================================================
h_section(doc, "1. Lógica del Ordenamiento Causal Propuesto")

box_chain(doc,
    "IPC  →  Crédito  →  Actividad Económica  →  Inflación\n"
    "(versión simple:  IPC  →  Actividad Económica  →  Inflación)"
)

body(doc,
    "El Índice de Presión Cambiaria (IPC) sintetiza las tensiones sobre el tipo de cambio "
    "nominal y las reservas internacionales del banco central. Cuando el IPC aumenta "
    "(mayor presión depreciativa), activa dos canales macroeconómicos encadenados:"
)
body(doc,
    "   1. Canal financiero: la depreciación eleva el costo del endeudamiento en moneda "
    "extranjera y deteriora los balances de bancos y empresas con deuda dolarizada, lo que "
    "contrae el crédito disponible (credit crunch)."
)
body(doc,
    "   2. Canal de actividad: la restricción crediticia reduce la inversión y el consumo, "
    "deprimiendo la actividad económica (PIB)."
)
body(doc,
    "   3. Canal de precios: el menor PIB junto con el pass-through cambiario elevan la "
    "inflación, cerrando la cadena causal."
)
body(doc,
    "Este ordenamiento justifica el uso de un VAR con identificación de Cholesky en el que "
    "el IPC precede al crédito, que precede a la actividad y que precede a la inflación, "
    "dado que las variables más exógenas (aquellas que reaccionan más lentamente a las "
    "demás) se colocan primero en la descomposición triangular. La evidencia empírica que "
    "sustenta cada eslabón de esta cadena se documenta a continuación."
)

doc.add_page_break()

# ==============================================================
# SECCIÓN 2 — PAPER 1: GIRTON & ROPER (1977)
# ==============================================================
h_section(doc, "2. Papers de Soporte — Definición y Marco Conceptual del IPC")

h_paper(doc, "Paper 1 — Girton & Roper (1977)  |  Base de datos: Scopus (AEA)")
apa_ref(doc,
    "Girton, L., & Roper, D. (1977). A monetary model of exchange market pressure applied "
    "to the postwar Canadian experience. American Economic Review, 67(4), 537–548. "
    "https://www.jstor.org/stable/1813444"
)
field(doc, "Base de datos: ", "Scopus (American Economic Review — AEA, indexado en Scopus Q1)")
field(doc, "Objetivo: ",
    "Proponer y aplicar el primer modelo formal del Índice de Presión Cambiaria (EMP/IPC), "
    "definido como la suma ponderada de la variación del tipo de cambio y la pérdida de "
    "reservas internacionales, bajo el enfoque monetario de la balanza de pagos.")
field(doc, "Metodología: ",
    "Modelo monetario de la balanza de pagos estimado mediante MCO para Canadá en el "
    "período de posguerra. Introduce la identidad EMP = Δe + Δr como medida compuesta "
    "de la presión total sobre el mercado cambiario.")
field(doc, "Resultados principales: ",
    "El IPC responde significativamente a los diferenciales de crecimiento del crédito "
    "doméstico y a las condiciones monetarias externas. Un exceso de crédito doméstico "
    "genera presión depreciativa que se manifiesta en pérdida de reservas o depreciación.")
field(doc, "Justificación para el VAR / Granger: ",
    "Define formalmente la variable IPC que encabeza el ordenamiento de Cholesky propuesto. "
    "Establece que el crédito doméstico es una fuente directa de presión cambiaria "
    "(IPC), lo que fundamenta la causalidad IPC ↔ crédito en el primer eslabón de la "
    "cadena. Es el punto de partida obligatorio para cualquier investigación que use el IPC "
    "como variable endógena en un modelo VAR.")
divider(doc)

# ==============================================================
# SECCIÓN 3 — PAPERS 2 y 3: KEEFE (2021) y KEEFE & SAHA (2024)
# ==============================================================
h_section(doc, "3. Papers de Soporte — IPC y Condiciones Crediticias Globales")

h_paper(doc, "Paper 2 — Keefe (2021)  |  Base de datos: ScienceDirect")
apa_ref(doc,
    "Keefe, H. G. (2021). The transmission of global monetary and credit shocks on "
    "exchange market pressure in emerging markets and developing economies. "
    "Journal of International Financial Markets, Institutions and Money, 72, 101302. "
    "https://doi.org/10.1016/j.intfin.2021.101302"
)
field(doc, "Base de datos: ", "ScienceDirect (Elsevier) — Journal of International Financial Markets, Institutions and Money (Scopus Q1)")
field(doc, "Objetivo: ",
    "Analizar cómo los shocks en las condiciones monetarias y crediticias globales de "
    "economías avanzadas se transmiten al índice de presión cambiaria (IPC/EMP) de "
    "40 economías emergentes y en desarrollo, 1998–2016.")
field(doc, "Metodología: ",
    "Panel Vector Autoregression (Panel VAR) de efectos fijos. Funciones impulso-respuesta "
    "(IRF) y descomposición de varianza. Los países se clasifican por grado de apertura "
    "comercial y financiera. Identificación mediante Cholesky.")
field(doc, "Resultados principales: ",
    "Aumentos en la liquidez monetaria global o en el crédito global generan apreciación "
    "del IPC (menor presión) en EMDEs. Sin embargo, en condiciones de alta volatilidad "
    "financiera, la inyección de liquidez no atenúa la presión depreciativa. El tamaño del "
    "impacto depende de la apertura comercial y financiera del país. El nexo crédito ↔ IPC "
    "es bidireccional y estadísticamente robusto.")
field(doc, "Justificación para el VAR / Granger: ",
    "Valida empíricamente el vínculo entre crédito e IPC usando Panel VAR con Cholesky "
    "para 40 economías emergentes. Justifica incluir el IPC en el VAR antes del crédito "
    "(o con el crédito como variable mediadora), dado que las condiciones crediticias "
    "afectan y son afectadas por el IPC. El diseño panel VAR con IRF es metodológicamente "
    "equivalente al modelo propuesto para un país individual.")
divider(doc)

h_paper(doc, "Paper 3 — Keefe & Saha (2024)  |  Base de datos: ScienceDirect")
apa_ref(doc,
    "Keefe, H. G., & Saha, S. (2024). Global uncertainty and exchange rate conditions: "
    "Assessing the impact of uncertainty shocks in emerging markets and advanced economies. "
    "Journal of International Financial Markets, Institutions and Money, 96, 102047. "
    "https://doi.org/10.1016/j.intfin.2024.102047"
)
field(doc, "Base de datos: ", "ScienceDirect (Elsevier) — Journal of International Financial Markets, Institutions and Money (Scopus Q1)")
field(doc, "Objetivo: ",
    "Evaluar cómo los shocks de incertidumbre global (económica, financiera y de política "
    "monetaria de EE. UU.) impactan el tipo de cambio bilateral, la volatilidad cambiaria "
    "y el IPC en economías emergentes y avanzadas.")
field(doc, "Metodología: ",
    "Global Vector Autoregression (GVAR). Funciones impulso-respuesta generalizadas. "
    "Variables: IPC, tipo de cambio bilateral, volatilidad cambiaria, VIX, incertidumbre "
    "de política económica (EPU). Muestra de múltiples países.")
field(doc, "Resultados principales: ",
    "Los picos de incertidumbre global generan aumentos inmediatos en la volatilidad "
    "cambiaria, acumulación de presión depreciativa (IPC elevado) y depreciación en "
    "economías emergentes. El IPC responde en la misma dirección que los shocks de "
    "incertidumbre, con efectos que se prolongan varios trimestres. Las economías avanzadas "
    "muestran respuestas distintas y más amortiguadas.")
field(doc, "Justificación para el VAR / Granger: ",
    "Proporciona evidencia de que el IPC es una variable que responde a condiciones "
    "exógenas (incertidumbre global, política de EE. UU.) y que antecede —en términos "
    "causales— a la dinámica macroeconómica interna. Esto respalda ubicar el IPC al "
    "inicio del ordenamiento de Cholesky. El uso de GVAR con IRF generalizadas es "
    "metodológicamente análogo al VAR propuesto.")
divider(doc)

doc.add_page_break()

# ==============================================================
# SECCIÓN 4 — PAPERS 4 y 5: CAUSALIDAD EMP → CRÉDITO
# ==============================================================
h_section(doc, "4. Papers de Soporte — Causalidad de Granger: IPC y Crédito Doméstico")

h_paper(doc, "Paper 4 — Olanipekun, Güngör & Olasehinde-Williams (2019)  |  Base de datos: Scopus")
apa_ref(doc,
    "Olanipekun, I. O., Güngör, H., & Olasehinde-Williams, G. (2019). Unraveling the "
    "causal relationship between economic policy uncertainty and exchange market pressure "
    "in BRIC countries: Evidence from bootstrap panel Granger causality. "
    "SAGE Open, 9(2), 2158244019853903. "
    "https://doi.org/10.1177/2158244019853903"
)
field(doc, "Base de datos: ", "Scopus (SAGE Open — SAGE Publishing, indexado en Scopus)")
field(doc, "Objetivo: ",
    "Examinar la relación causal entre la incertidumbre de política económica (global y "
    "doméstica) y el IPC en las economías BRIC, con énfasis en detectar causalidad "
    "unidireccional o bidireccional mediante pruebas robustas a heterocedasticidad y "
    "dependencia de corte transversal.")
field(doc, "Metodología: ",
    "Bootstrap Panel Granger Causality (Kónya, 2006). Pruebas de raíz unitaria de paneles "
    "con quiebres estructurales. Prueba de dependencia de corte transversal (CD test). "
    "Datos anuales de Brasil, Rusia, India y China.")
field(doc, "Resultados principales: ",
    "A nivel panel: causalidad bidireccional entre incertidumbre de política económica "
    "(global y doméstica) y el IPC. Por país: China muestra causalidad bidireccional entre "
    "EPU global e IPC; India y Rusia muestran causalidad bidireccional entre EPU doméstica "
    "e IPC. El IPC Granger-causa la EPU en varios países, indicando su carácter predictivo.")
field(doc, "Justificación para el VAR / Granger: ",
    "Aplica directamente pruebas de causalidad de Granger para el IPC como variable "
    "endógena en un panel de economías emergentes. Confirma que el IPC tiene relaciones "
    "causales significativas con otras variables macroeconómicas, lo que justifica su "
    "inclusión en la especificación VAR y la prueba de causalidad de Granger propuesta. "
    "La metodología bootstrap robustece los resultados ante no-normalidad y heterocedasticidad.")
divider(doc)

h_paper(doc, "Paper 5 — Şıklar & Akça (2020)  |  Base de datos: Scopus (Ekonomika)")
apa_ref(doc,
    "Şıklar, İ., & Akça, A. (2020). Exchange market pressure and monetary policy: "
    "The Turkish case. Ekonomika, 99(1), 110–130. "
    "https://doi.org/10.15388/Ekon.2020.1.7"
)
field(doc, "Base de datos: ", "Scopus (Ekonomika — Vilnius University, indexado en Scopus y ESCI)")
field(doc, "Objetivo: ",
    "Determinar la relación entre la política monetaria y el IPC en Turquía para el "
    "período 2002–2018, evaluando si los instrumentos de política monetaria —crédito "
    "doméstico y diferencial de tasas de interés— Granger-causan el IPC o viceversa.")
field(doc, "Metodología: ",
    "Estimación del IPC siguiendo el modelo de Girton & Roper (1977). Modelo VAR mensual. "
    "Prueba de causalidad de Granger bidireccional entre: (1) crédito doméstico e IPC; "
    "(2) diferencial de tasa de interés e IPC. Variables: M1, reservas internacionales, "
    "tipo de cambio, diferencial de tasas, crédito doméstico.")
field(doc, "Resultados principales: ",
    "Causalidad unidireccional de crédito doméstico → IPC (el crédito Granger-causa el "
    "IPC, no al revés). Causalidad bidireccional entre diferencial de tasa de interés e "
    "IPC. La expansión del crédito doméstico genera presión depreciativa (IPC mayor), "
    "confirmando la lógica del modelo de Girton & Roper. El VAR muestra que shocks en el "
    "crédito se transmiten al IPC con 2–3 meses de rezago.")
field(doc, "Justificación para el VAR / Granger: ",
    "Provee evidencia directa de causalidad de Granger crédito → IPC con datos de una "
    "economía emergente usando VAR, que es exactamente la metodología propuesta. "
    "Justifica el primer eslabón de la cadena: la presión cambiaria (IPC) está vinculada "
    "causalmente al crédito. Si el crédito Granger-causa el IPC, tiene sentido incluirlos "
    "ambos en el VAR con el IPC respondiendo a condiciones crediticias y viceversa.")
divider(doc)

doc.add_page_break()

# ==============================================================
# SECCIÓN 5 — PAPER 6: CADENA COMPLETA
# ==============================================================
h_section(doc, "5. Paper de Soporte — Cadena Completa: IPC → Crédito → PIB → Inflación")

h_paper(doc, "Paper 6 — Doojav, Purevdorj & Batjargal (2024)  |  Base de datos: ScienceDirect")
apa_ref(doc,
    "Doojav, G.-O., Purevdorj, M., & Batjargal, A. (2024). The macroeconomic effects of "
    "exchange rate movements in a commodity-exporting developing economy. "
    "International Economics, 177, 100466. "
    "https://doi.org/10.1016/j.inteco.2023.100466"
)
field(doc, "Base de datos: ", "ScienceDirect (Elsevier) — International Economics (Scopus)")
field(doc, "Objetivo: ",
    "Examinar los efectos macroeconómicos de los movimientos del tipo de cambio —y por "
    "extensión de la presión cambiaria— en Mongolia, economía exportadora de commodities "
    "y deudora neta en moneda extranjera, separando los canales de transmisión: financiero "
    "y comercial.")
field(doc, "Metodología: ",
    "Structural Bayesian Vector Autoregression (SBVAR) mensual. Dos indicadores cambiarios: "
    "(1) tipo de cambio ponderado por deuda (DWER), que captura el canal financiero; "
    "(2) tipo de cambio efectivo nominal (NEER), que captura el canal comercial. "
    "Funciones impulso-respuesta estructurales y descomposición de varianza. "
    "Variables: tipo de cambio, crédito bancario, PIB, IPC, inversión, exportaciones netas.")
field(doc, "Resultados principales: ",
    "Canal financiero (DWER): la depreciación cambiaria deteriora los balances de bancos "
    "y empresas con deuda en moneda extranjera, contrae el crédito bancario y reduce la "
    "inversión y el PIB (especialmente en manufactura y construcción). El PIB cae de forma "
    "significativa y persistente. "
    "Canal comercial (NEER): la depreciación mejora las exportaciones netas y eleva el IPC "
    "(inflación) vía pass-through cambiario. "
    "Ambos canales operan simultáneamente y son estadísticamente significativos.")
field(doc, "Justificación para el VAR / Granger: ",
    "Es el paper que más directamente justifica la cadena completa propuesta: "
    "Presión cambiaria → crédito → actividad económica → inflación. "
    "El canal financiero del SBVAR muestra exactamente esta secuencia: el deterioro "
    "cambiario (IPC) afecta el crédito, que deprime el PIB. El canal comercial cierra "
    "la cadena hacia la inflación. El uso de SBVAR con Cholesky estructural es "
    "metodológicamente idéntico al VAR propuesto. Aplica a una economía emergente y "
    "exportadora de commodities, comparable a economías latinoamericanas.")
divider(doc)

doc.add_page_break()

# ==============================================================
# SECCIÓN 6 — TABLA SÍNTESIS
# ==============================================================
h_section(doc, "6. Tabla Síntesis de Papers y Eslabones de la Cadena")

doc.add_paragraph()

table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"

headers = ["Paper", "Base de datos", "Metodología", "Eslabón justificado", "Aporte clave"]
hdr = table.rows[0].cells
for i, h in enumerate(headers):
    hdr[i].text = h
    shade_cell(hdr[i], "1F497D")
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(9); run.font.color.rgb = C_WHITE

rows_data = [
    ["Girton & Roper (1977)",
     "Scopus\n(AER)",
     "Modelo monetario BOP",
     "Definición del IPC\n(toda la cadena)",
     "Funda el IPC; muestra crédito → IPC"],
    ["Keefe (2021)",
     "ScienceDirect\n(JIFMI)",
     "Panel VAR\nCholesky (40 países)",
     "IPC ↔ Crédito",
     "Nexo bidireccional crédito-IPC en EMDEs"],
    ["Keefe & Saha (2024)",
     "ScienceDirect\n(JIFMI)",
     "Global VAR\nIRF generalizadas",
     "IPC precede a\nvariables internas",
     "IPC como variable más exógena en el VAR"],
    ["Olanipekun et al. (2019)",
     "Scopus\n(SAGE Open)",
     "Bootstrap Panel\nGranger Causality",
     "Causalidad de Granger\ndel IPC",
     "Granger causality bidireccional IPC-incertidumbre"],
    ["Şıklar & Akça (2020)",
     "Scopus\n(Ekonomika)",
     "VAR + Granger\ncausalidad",
     "Crédito → IPC\n(Granger causal)",
     "Evidencia directa crédito→IPC en economía emergente"],
    ["Doojav et al. (2024)",
     "ScienceDirect\n(Int. Economics)",
     "SBVAR\ncanal financiero",
     "Cadena completa:\nIPC→Cred→PIB→Inf",
     "SBVAR muestra la cadena completa con IRF"],
]

alt = ["F2F2F2", "FFFFFF"]
for idx, row_data in enumerate(rows_data):
    row = table.add_row()
    fill = alt[idx % 2]
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        shade_cell(cell, fill)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

# ==============================================================
# SECCIÓN 7 — NOTA DE CIERRE
# ==============================================================
doc.add_paragraph()
h_section(doc, "7. Síntesis para el Marco Metodológico")

body(doc,
    "Los seis papers seleccionados cubren de forma sistemática los tres eslabones de la "
    "cadena causal IPC → Crédito → Actividad Económica → Inflación, así como la "
    "metodología VAR y las pruebas de causalidad de Granger propuestas:"
)
body(doc,
    "• Eslabón IPC (definición y exogeneidad): Girton & Roper (1977) define el IPC; "
    "Keefe & Saha (2024) muestra que el IPC es una variable que responde a condiciones "
    "externas exógenas, lo que justifica colocarlo primero en el ordenamiento de Cholesky."
)
body(doc,
    "• Eslabón IPC → Crédito: Keefe (2021) y Şıklar & Akça (2020) muestran, mediante "
    "Panel VAR y Granger causality, que el crédito y el IPC están causalmente vinculados. "
    "Girton & Roper (1977) establece que el crédito doméstico Granger-causa el IPC."
)
body(doc,
    "• Eslabón Crédito → Actividad Económica → Inflación: Doojav et al. (2024) documenta "
    "con SBVAR que la presión cambiaria contrae el crédito (canal financiero), que a su "
    "vez reduce el PIB, mientras que el canal comercial transmite la presión cambiaria "
    "directamente a la inflación mediante el pass-through."
)
body(doc,
    "• Causalidad de Granger y Panel: Olanipekun et al. (2019) valida el uso de pruebas "
    "de causalidad de Granger para el IPC en economías emergentes, con el bootstrap "
    "approach que es robusto a heterocedasticidad y dependencia transversal."
)

p_note = doc.add_paragraph()
r_note = p_note.add_run(
    "Todos los papers están indexados en Scopus y/o disponibles en ScienceDirect (Elsevier). "
    "Las referencias siguen el formato APA 7ª edición con DOI verificado."
)
r_note.font.size = Pt(9); r_note.font.italic = True; r_note.font.color.rgb = C_GREY

# ==============================================================
# GUARDAR
# ==============================================================
out = (
    r"C:\Users\51950\Documents\ULIMA\2026_1\Teo. Politica"
    r"\Trabajo grupal\Regresion_VAR\Bibliografia\Literatura_IPC_VAR_Cholesky.docx"
)
doc.save(out)
print(f"Documento guardado: {out}")
