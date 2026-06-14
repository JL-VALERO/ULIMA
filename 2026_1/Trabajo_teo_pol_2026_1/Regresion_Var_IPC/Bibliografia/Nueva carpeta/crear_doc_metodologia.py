"""
Genera Metodologia_VAR_IPC.docx con todas las notas metodologicas
para el modelo VAR: IPC -> Credito -> Actividad -> Inflacion
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

AZUL     = RGBColor(0x1F, 0x49, 0x7D)
VERDE    = RGBColor(0x37, 0x58, 0x23)
GRIS     = RGBColor(0x40, 0x40, 0x40)
NARANJA  = RGBColor(0xC0, 0x50, 0x00)

def set_run_fmt(run, bold=False, italic=False, size=11, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_fmt(run, bold=True, size=14, color=AZUL)

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_run_fmt(run, bold=True, size=12, color=VERDE)

def body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_fmt(run, italic=italic, size=11)

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.4)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = NARANJA

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ": ")
        set_run_fmt(r1, bold=True, size=11)
    r2 = p.add_run(text)
    set_run_fmt(r2, size=11)

def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # Header
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for par in hdr_cells[i].paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F497D")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        cells = tbl.rows[ri+1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            for par in cells[ci].paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

# ============================================================
doc = Document()

# --- Titulo ---
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(4)
r = t.add_run("Guia Metodologica: Modelo VAR con Indice de Presion Cambiaria")
set_run_fmt(r, bold=True, size=16, color=AZUL)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(12)
rs = sub.add_run("Teoria y Politica Monetaria — ULIMA 2026-1")
set_run_fmt(rs, italic=True, size=12, color=GRIS)

# ============================================================
heading1(doc, "1. El Modelo Propuesto")
body(doc, "Propuesta 1 seleccionada: ordenamiento Cholesky IPC -> Credito -> "
          "Actividad Economica -> Inflacion.")
body(doc, "El Indice de Presion Cambiaria (IPC) = delta(tipo de cambio) + delta(reservas) "
          "siguiendo Girton & Roper (1977). Captura simultaneamente la presion de mercado "
          "sobre el sol Y la respuesta del BCRP via intervencion en reservas.")

heading2(doc, "1.1 Justificacion del ordenamiento Cholesky")
bullet(doc, "Responde solo a shocks externos. No reacciona contemporaneamente al credito "
            "ni a la inflacion domestica. Es la variable mas exogena del sistema.",
       bold_prefix="IPC (1ro)")
bullet(doc, "Los bancos ajustan portafolios ante presion cambiaria antes de que el PBI "
            "lo registre. Dollarizacion: balance-sheet effect (Cespedes, Chang & Velasco 2004).",
       bold_prefix="Credito (2do)")
bullet(doc, "El PBI mensual es el resultado rezagado de condiciones financieras. "
            "Tarda en reaccionar al endurecimiento crediticio.",
       bold_prefix="Actividad (3ro)")
bullet(doc, "La variable mas endogena. Responde a todas las anteriores: presion cambiaria "
            "via pass-through + actividad via brecha del producto.",
       bold_prefix="Inflacion (4to)")

# ============================================================
heading1(doc, "2. Variables y Fuentes de Datos")

add_table(doc,
    ["Variable", "Serie BCRP", "Codigo", "Unidad", "Transformacion VAR"],
    [
        ["IPC (Presion Cambiaria)", "Construido por el autor", "—", "Indice / %", "Niveles si I(0)"],
        ["Produccion Nacional", "Prod. Nacional Indice 2007=100", "PN01770AM", "Indice", "Dln x 100"],
        ["Credito Privado", "Credito Sist. Financiero", "PN00518MM", "Mill. S/", "Dln x 100"],
        ["Credito Publico (neto)", "Credito neto sector publico", "PN00881MM", "Mill. S/", "Nivel o Dif."],
        ["Tasa Interbancaria", "Tasa interbancaria prom. MN", "PN07819NM", "% TEA", "Nivel directo"],
        ["M1 (Masa monetaria)", "Dinero Sist. Financiero", "PN00199MM", "Mill. S/", "Dln x 100"],
    ]
)
body(doc, "Archivo Excel de datos: datos/datos_var_ipc.xlsx (264 obs, Ene 2004 – Dic 2025, 0 faltantes).")

# ============================================================
heading1(doc, "3. Transformacion de Variables")

heading2(doc, "3.1 IPC — Indice de Presion Cambiaria")
body(doc, "Ya es una diferencia de variables de nivel (delta TC + delta Reservas). "
          "Es probable que sea I(0). Verificar con ADF antes de incluir en el VAR.")
body(doc, "Si resulta I(0): usar directamente en niveles (puntos porcentuales).")
body(doc, "Si resulta I(1): aplicar una diferencia adicional.")

heading2(doc, "3.2 Credito al Sector Privado")
body(doc, "El nivel en millones de soles tiene tendencia fuerte -> I(1) o I(2). "
          "Aplicar log-transformacion y luego primera diferencia en Stata.")
code_block(doc, "gen ln_credito = ln(credito_privado)")
code_block(doc, "* En el VAR usar: D.ln_credito  (tasa de crecimiento mensual)")
body(doc, "Nota: usar credito nominal. Como la inflacion ya esta en el VAR, "
          "el modelo descompone el efecto real del monetario endogenamente. "
          "Deflactar crearia correlacion mecanica artificial entre credito e inflacion.")

heading2(doc, "3.3 Actividad Economica (Produccion Nacional)")
body(doc, "El indice de produccion tiene tendencia + estacionalidad -> I(1) con componente estacional.")
code_block(doc, "gen ln_prod = ln(produccion_nacional)")
code_block(doc, "* En el VAR usar: D.ln_prod")
body(doc, "IMPORTANTE: La produccion mensual tiene estacionalidad fuerte (diciembre alto, enero bajo). "
          "Dos opciones:")
bullet(doc, "Usar variacion año a año: S12.ln_prod = ln(prod_t) - ln(prod_{t-12}). "
            "Simple pero introduce autocorrelacion MA(11) en errores.")
bullet(doc, "Usar la serie desestacionalizada del INEI directamente. RECOMENDADO.")
bullet(doc, "Incluir dummies estacionales mensuales en el VAR como variables exogenas. "
            "Consume 11 grados de libertad.")

heading2(doc, "3.4 Inflacion")
body(doc, "La tasa de inflacion mensual (variacion del IPC precios al consumidor) "
          "ya es una tasa de cambio -> tipicamente I(0).")
body(doc, "Usar directamente en porcentaje. No diferenciar. Verificar con ADF.")
body(doc, "DISTINGUIR: inflacion (variacion del IPC precios) != IPC presion cambiaria. "
          "Son variables distintas con nombres similares.")

heading2(doc, "3.5 Tasa Interbancaria")
body(doc, "Ya es una tasa en %. Usar directamente en niveles. "
          "Suele ser I(1) o I(0) dependiendo del periodo. Verificar con ADF.")
body(doc, "Si resulta I(1): diferenciar o considerar que el VAR en niveles sigue siendo "
          "valido asintoticamente (Sims, Stock & Watson 1990).")

heading2(doc, "3.6 M1 — Masa Monetaria")
body(doc, "El nivel en millones de soles tiene tendencia fuerte -> I(1). "
          "Aplicar log + primera diferencia para obtener tasa de crecimiento de M1.")
code_block(doc, "gen ln_m1 = ln(m1_dinero)")
code_block(doc, "* En el VAR usar: D.ln_m1")

# ============================================================
heading1(doc, "4. Protocolo de Pruebas de Raiz Unitaria en Stata")

body(doc, "Correr para CADA variable transformada antes de estimar el VAR:")
code_block(doc, "* ADF con constante + tendencia")
code_block(doc, "dfuller ipc_presion,    lags(4) trend")
code_block(doc, "dfuller D.ln_credito,   lags(4) trend")
code_block(doc, "dfuller D.ln_prod,      lags(4) trend")
code_block(doc, "dfuller inflacion,      lags(4) trend")
code_block(doc, "dfuller tasa_interbancaria, lags(4) trend")
code_block(doc, "dfuller D.ln_m1,        lags(4) trend")
code_block(doc, "")
code_block(doc, "* Zivot-Andrews: permite quiebre estructural (recomendado con GFC y COVID)")
code_block(doc, "zandrews ipc_presion, maxlags(4)")
code_block(doc, "zandrews D.ln_credito, maxlags(4)")

body(doc, "Hipotesis H0 en ADF: tiene raiz unitaria (es I(1)). "
          "Si p-value < 0.05: rechazo H0 -> variable es I(0) -> OK para VAR en niveles.")
body(doc, "Si alguna variable resulta I(1) despues de las transformaciones propuestas: "
          "aplicar una diferencia adicional. Si todas son I(0): estimar VAR directamente.")

# ============================================================
heading1(doc, "5. Codigo Stata — Estructura del VAR")

heading2(doc, "5.1 Importar datos y preparar variables")
code_block(doc, "import excel using \"datos/datos_var_ipc.xlsx\", sheet(\"niveles\") firstrow clear")
code_block(doc, "gen fecha = ym(year(periodo), month(periodo))")
code_block(doc, "format fecha %tm")
code_block(doc, "tsset fecha")
code_block(doc, "")
code_block(doc, "* Transformaciones")
code_block(doc, "gen ln_prod    = ln(produccion_nacional)")
code_block(doc, "gen ln_cred    = ln(credito_privado)")
code_block(doc, "gen ln_m1      = ln(m1_dinero)")
code_block(doc, "* Primera diferencia")
code_block(doc, "gen g_prod     = D.ln_prod * 100   // crecimiento mensual actividad")
code_block(doc, "gen g_cred     = D.ln_cred * 100   // crecimiento mensual credito")
code_block(doc, "gen g_m1       = D.ln_m1  * 100   // crecimiento mensual M1")

heading2(doc, "5.2 Seleccion de rezagos optimos")
code_block(doc, "varsoc ipc g_cred g_prod inflacion, maxlag(6)")
code_block(doc, "* Elegir rezago segun AIC/HQIC/SBIC. Generalmente 1 o 2 para datos mensuales.")

heading2(doc, "5.3 Estimacion del VAR")
code_block(doc, "var ipc g_cred g_prod inflacion, lags(1/2)")
code_block(doc, "* Orden Cholesky: IPC -> Credito -> Actividad -> Inflacion")

heading2(doc, "5.4 Diagnosticos")
code_block(doc, "varstable           // Estabilidad: eigenvalores < 1")
code_block(doc, "varlmar, mlag(4)    // Test LM autocorrelacion (H0: no autocorr)")
code_block(doc, "varnorm             // Normalidad Jarque-Bera")
code_block(doc, "vargranger          // Causalidad de Granger")

heading2(doc, "5.5 Funciones de Impulso-Respuesta (IRF)")
code_block(doc, "irf create irf_ipc, step(24) set(\"output/irf_ipc\") replace")
code_block(doc, "irf graph oirf, impulse(ipc) response(g_cred g_prod inflacion) ///")
code_block(doc, "    yline(0) scheme(s2color) name(irf_ipc, replace)")
code_block(doc, "graph export \"output/irf_ipc_respuestas.png\", replace")

heading2(doc, "5.6 Descomposicion de varianza (FEVD)")
code_block(doc, "irf graph fevd, impulse(ipc) response(g_cred g_prod inflacion) ///")
code_block(doc, "    yline(0) scheme(s2color)")

# ============================================================
heading1(doc, "6. Dummies Estructurales Recomendadas")

add_table(doc,
    ["Dummy", "Periodo", "Evento", "Justificacion"],
    [
        ["D_gfc", "Sep 2008 – Jun 2009", "Crisis Financiera Global", "Caida PBI, intervencion BCRP masiva"],
        ["D_pandemic", "Mar 2020 – Dic 2021", "Pandemia COVID-19", "Mayor distorsion de series en historia reciente"],
        ["D_highinfl", "Ene 2022 – Dic 2023", "Inflacion alta / post-COVID", "BCRP sube tasa 0.25% a 7.75%"],
    ]
)
code_block(doc, "gen D_gfc      = (fecha >= tm(2008m9)  & fecha <= tm(2009m6))")
code_block(doc, "gen D_pandemic = (fecha >= tm(2020m3)  & fecha <= tm(2021m12))")
code_block(doc, "gen D_highinfl = (fecha >= tm(2022m1)  & fecha <= tm(2023m12))")
code_block(doc, "")
code_block(doc, "* VAR con dummies como exogenas:")
code_block(doc, "var ipc g_cred g_prod inflacion, lags(1/2) exog(D_gfc D_pandemic D_highinfl)")

# ============================================================
heading1(doc, "7. Resumen de Codigos BCRP")

add_table(doc,
    ["Variable", "Codigo BCRP", "URL BCRP"],
    [
        ["Produccion Nacional (Indice 2007=100)", "PN01770AM",
         "estadisticas.bcrp.gob.pe/.../PN01770AM"],
        ["Credito Sector Privado (Mill. S/)", "PN00518MM",
         "estadisticas.bcrp.gob.pe/.../PN00518MM"],
        ["Credito Neto Sector Publico (Mill. S/)", "PN00881MM",
         "estadisticas.bcrp.gob.pe/.../PN00881MM"],
        ["Tasa Interbancaria Prom. MN (% TEA)", "PN07819NM",
         "estadisticas.bcrp.gob.pe/.../PN07819NM"],
        ["M1 Dinero Sist. Financiero (Mill. S/)", "PN00199MM",
         "estadisticas.bcrp.gob.pe/.../PN00199MM"],
    ]
)
body(doc, "Datos descargados automaticamente via script Python (descargar_series_bcrp.py). "
          "Para actualizar: re-ejecutar el script. Periodo: Ene 2004 – Dic 2025, 264 obs mensuales.")

# --- Guardar ---
out = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Metodologia_VAR_IPC.docx"
)
doc.save(out)
print(f"Documento guardado: {out}")
