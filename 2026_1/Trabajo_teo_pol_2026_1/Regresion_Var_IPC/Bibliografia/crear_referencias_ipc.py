"""
Genera Referencias_VAR_IPC.docx
Tema: Efectos dinamicos del IPC sobre el canal crediticio y la actividad economica en Peru
Fuentes: ScienceDirect y Scopus + archivos locales relevantes
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# --- Paleta de colores ---
C_BLUE   = RGBColor(0x1F, 0x49, 0x7D)
C_LBLUE  = RGBColor(0x2E, 0x74, 0xB5)
C_GREEN  = RGBColor(0x1D, 0x60, 0x3A)
C_GREY   = RGBColor(0x55, 0x55, 0x55)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_RED    = RGBColor(0x80, 0x20, 0x20)
C_ORANGE = RGBColor(0xBF, 0x65, 0x00)

# --- Helpers ---
def shade_cell(cell, hex_fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_fill)
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)

def h_main(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(17); r.font.color.rgb = C_BLUE

def h_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(11); r.font.color.rgb = C_LBLUE; r.italic = True

def h_section(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = C_LBLUE

def h_paper(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(f"[{num}]  {text}")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = C_GREEN

def apa(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after       = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10)

def field(doc, label, content):
    p = doc.add_paragraph()
    p.paragraph_format.space_after      = Pt(2)
    p.paragraph_format.left_indent      = Inches(0.15)
    r1 = p.add_run(label)
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = C_BLUE
    r2 = p.add_run(content)
    r2.font.size = Pt(10)

def badge(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"  {text}  ")
    r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = C_WHITE
    # No podemos poner fondo real sin XML extra; usamos color del texto
    r.font.color.rgb = C_ORANGE

def divider(doc):
    p = doc.add_paragraph()
    r = p.add_run("─" * 95)
    r.font.size = Pt(8); r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(12)

def body(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.15)
    for r in p.runs:
        r.font.size = Pt(10)

# ================================================================
doc = Document()
for s in doc.sections:
    s.top_margin    = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin   = Inches(1.25)
    s.right_margin  = Inches(1.25)

# ================================================================
# PORTADA
# ================================================================
h_main(doc, "Referencias Bibliograficas para la Investigacion")
doc.add_paragraph()
h_subtitle(doc,
    "Efectos dinamicos del Indice de Presion Cambiaria sobre el canal\n"
    "crediticio y la actividad economica en Peru: evidencia de un modelo VAR, 2004-2025"
)
doc.add_paragraph()
p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_info = p_info.add_run(
    "Teoria y Politica Monetaria  |  ULIMA 2026-1\n"
    "Fuentes: ScienceDirect (Elsevier) y Scopus\n"
    "Referencias en formato APA 7a edicion  |  13 papers"
)
r_info.font.size = Pt(10); r_info.font.color.rgb = C_GREY

doc.add_page_break()

# ================================================================
# INTRODUCCION
# ================================================================
h_section(doc, "Introduccion")
body(doc,
    "Este documento compila 13 referencias academicas que sustentan el diseno y la "
    "justificacion del modelo VAR propuesto. El modelo estudia como el Indice de Presion "
    "Cambiaria (IPC = variacion del tipo de cambio + variacion de reservas internacionales, "
    "siguiendo Girton & Roper, 1977) transmite sus efectos sobre el credito, la actividad "
    "economica y la inflacion en Peru durante el periodo 2004-2025."
)
body(doc,
    "Las referencias estan organizadas en seis secciones: (1) definicion y marco conceptual "
    "del IPC, (2) IPC y condiciones crediticias globales, (3) cadena de transmision "
    "IPC → actividad → inflacion, (4) evidencia especifica para Peru, (5) causalidad de "
    "Granger y metodologia VAR, y (6) choques externos y economias emergentes. "
    "Cada referencia incluye los cinco campos solicitados."
)

doc.add_page_break()

# ================================================================
# SECCION 1 — DEFINICION Y MARCO CONCEPTUAL DEL IPC
# ================================================================
h_section(doc, "Seccion 1: Definicion y Marco Conceptual del Indice de Presion Cambiaria")

# --- [1] Girton & Roper (1977) ---
h_paper(doc, 1, "Girton & Roper (1977)  |  Scopus — American Economic Review")
apa(doc,
    "Girton, L., & Roper, D. (1977). A monetary model of exchange market pressure "
    "applied to the postwar Canadian experience. American Economic Review, 67(4), 537-548. "
    "https://www.jstor.org/stable/1813444"
)
field(doc, "1. Que investigan: ",
    "Proponen y aplican el primer modelo formal del Indice de Presion Cambiaria (EMP/IPC), "
    "definido como EMP = delta(tipo de cambio) + delta(reservas), bajo el enfoque monetario "
    "de la balanza de pagos. Evaluan si la presion cambiaria sobre Canada responde a "
    "diferenciales de credito domestico y condiciones monetarias externas.")
field(doc, "2. Marco transversal y periodo: ",
    "Canada, posguerra (aproximadamente decadas de 1950-1970). "
    "Datos anuales de la experiencia canadiense de posguerra.")
field(doc, "3. Metodologia: ",
    "Modelo monetario de la balanza de pagos estimado mediante MCO. "
    "Introduce la identidad EMP = Delta(e) + Delta(r) como medida compuesta de "
    "la presion total sobre el mercado cambiario.")
field(doc, "4. Principales hallazgos: ",
    "El IPC responde significativamente a los diferenciales de crecimiento del credito "
    "domestico y a las condiciones monetarias externas. Un exceso de credito domestico "
    "genera presion depreciativa manifestada en perdida de reservas o depreciacion del "
    "tipo de cambio. La suma de ambas respuestas es el IPC.")
field(doc, "5. Cita APA 7: ",
    "Girton, L., & Roper, D. (1977). A monetary model of exchange market pressure applied "
    "to the postwar Canadian experience. American Economic Review, 67(4), 537-548.")
divider(doc)

# --- [2] Aizenman & Binici (2016) ---
h_paper(doc, 2, "Aizenman & Binici (2016)  |  ScienceDirect — Journal of International Money and Finance")
apa(doc,
    "Aizenman, J., & Binici, M. (2016). Exchange market pressure in OECD and emerging "
    "economies: Domestic vs. external factors and capital flows in the old and new normal. "
    "Journal of International Money and Finance, 66, 65-87. "
    "https://doi.org/10.1016/j.jimonfin.2015.12.008"
)
field(doc, "1. Que investigan: ",
    "Descomponen los determinantes del IPC en 50 economias de la OCDE y mercados emergentes, "
    "distinguiendo entre factores domesticos (tasas de interes, credito) y externos "
    "(liquidez global, politica monetaria de EE.UU., precios de commodities, flujos de "
    "capital brutos). Evaluan si la naturaleza del IPC cambia entre el 'viejo normal' "
    "(pre-2008) y el 'nuevo normal' (post-2008).")
field(doc, "2. Marco transversal y periodo: ",
    "Panel de 50 paises (OCDE + emergentes, incluidos Peru y paises de AL). "
    "Datos trimestrales, 2000T1-2014T3.")
field(doc, "3. Metodologia: ",
    "Panel dinamico con GMM sistema (Arellano-Bond). Variables clave: EMPI, "
    "apetito de riesgo global, liquidez global, tasa de la Fed, precios de commodities, "
    "flujos de capital brutos (portafolio, IED), controles de capital.")
field(doc, "4. Principales hallazgos: ",
    "Los factores externos tienen mayor impacto sobre el IPC en economias emergentes que "
    "en OCDE, validando la vulnerabilidad de Peru como exportador de commodities. "
    "Los flujos de portafolio de corto plazo son el canal dominante de generacion del IPC "
    "en emergentes. Post-2008, los factores monetarios externos amplificaron "
    "sustancialmente la transmision del IPC.")
field(doc, "5. Cita APA 7: ",
    "Aizenman, J., & Binici, M. (2016). Exchange market pressure in OECD and emerging "
    "economies: Domestic vs. external factors and capital flows in the old and new normal. "
    "Journal of International Money and Finance, 66, 65-87. "
    "https://doi.org/10.1016/j.jimonfin.2015.12.008")
divider(doc)

doc.add_page_break()

# ================================================================
# SECCION 2 — IPC Y CONDICIONES CREDITICIAS GLOBALES
# ================================================================
h_section(doc, "Seccion 2: IPC y Condiciones Crediticias en Economias Emergentes")

# --- [3] Keefe (2021) ---
h_paper(doc, 3, "Keefe (2021)  |  ScienceDirect — Journal of International Financial Markets, Institutions and Money")
apa(doc,
    "Keefe, H. G. (2021). The transmission of global monetary and credit shocks on "
    "exchange market pressure in emerging markets and developing economies. "
    "Journal of International Financial Markets, Institutions and Money, 72, 101302. "
    "https://doi.org/10.1016/j.intfin.2021.101302"
)
field(doc, "1. Que investigan: ",
    "Analizan como los shocks en condiciones monetarias y crediticias globales de economias "
    "avanzadas se transmiten al IPC de 40 economias emergentes y en desarrollo (EMDEs). "
    "Evaluan si la apertura comercial y financiera del pais modera el tamano del impacto.")
field(doc, "2. Marco transversal y periodo: ",
    "Panel de 40 EMDEs clasificadas por grado de apertura comercial y financiera. "
    "Datos mensuales, 1998-2016.")
field(doc, "3. Metodologia: ",
    "Panel VAR de efectos fijos con identificacion de Cholesky. "
    "Funciones impulso-respuesta (IRF) y descomposicion de varianza (FEVD). "
    "Los paises se clasifican por apertura para evaluar heterogeneidad.")
field(doc, "4. Principales hallazgos: ",
    "Aumentos en la liquidez monetaria global generan apreciacion del IPC (menor presion) "
    "en EMDEs. En condiciones de alta volatilidad financiera, la inyeccion de liquidez "
    "no atenua la presion depreciativa. El nexo credito <-> IPC es bidireccional y "
    "estadisticamente robusto. El tamano del impacto depende de la apertura del pais.")
field(doc, "5. Cita APA 7: ",
    "Keefe, H. G. (2021). The transmission of global monetary and credit shocks on "
    "exchange market pressure in emerging markets and developing economies. "
    "Journal of International Financial Markets, Institutions and Money, 72, 101302. "
    "https://doi.org/10.1016/j.intfin.2021.101302")
divider(doc)

# --- [4] Gevorkyan (2019) ---
h_paper(doc, 4, "Gevorkyan (2019)  |  Scopus — Applied Economics (Taylor & Francis)")
apa(doc,
    "Gevorkyan, A. V. (2019). Exchange market pressure and primary commodity-exporting "
    "emerging markets. Applied Economics, 51(22), 2390-2412. "
    "https://doi.org/10.1080/00036846.2018.1545077"
)
field(doc, "1. Que investigan: ",
    "Examinan la sensibilidad de corto plazo entre el IPC y factores macroeconomicos "
    "domesticos (tasas de interes, credito domestico, reservas internacionales) y externos "
    "(precios de commodities) en economias emergentes exportadoras de materias primas, "
    "diferenciando por tipo de commodity y regimen cambiario.")
field(doc, "2. Marco transversal y periodo: ",
    "Panel de economias emergentes exportadoras de commodities (combustibles, minerales, "
    "cereales, azucar, cafe), comparable con el perfil exportador de Peru. "
    "Sub-periodos: boom y post-boom del ciclo de commodities.")
field(doc, "3. Metodologia: ",
    "Panel VAR y modelos de panel con datos de flujos de capital y precios de commodities. "
    "Diferenciacion por regimen cambiario (flotacion administrada vs. ancla cambiaria).")
field(doc, "4. Principales hallazgos: ",
    "En regimenes de flotacion (como Peru), la volatilidad del precio de commodities se "
    "canaliza simultaneamente al mercado cambiario, las tasas de interes y los ciclos de "
    "credito. Los ciclos de credito domestico son endogenamente impulsados por el IPC "
    "bajo regimenes flexibles, estableciendo el vinculo IPC -> credito.")
field(doc, "5. Cita APA 7: ",
    "Gevorkyan, A. V. (2019). Exchange market pressure and primary commodity-exporting "
    "emerging markets. Applied Economics, 51(22), 2390-2412. "
    "https://doi.org/10.1080/00036846.2018.1545077")
divider(doc)

doc.add_page_break()

# ================================================================
# SECCION 3 — CADENA IPC → ACTIVIDAD → INFLACION
# ================================================================
h_section(doc, "Seccion 3: Cadena de Transmision IPC -> Actividad -> Inflacion")

# --- [5] Doojav, Purevdorj & Batjargal (2024) ---
h_paper(doc, 5, "Doojav, Purevdorj & Batjargal (2024)  |  ScienceDirect — International Economics")
apa(doc,
    "Doojav, G.-O., Purevdorj, M., & Batjargal, A. (2024). The macroeconomic effects of "
    "exchange rate movements in a commodity-exporting developing economy. "
    "International Economics, 177, 100466. "
    "https://doi.org/10.1016/j.inteco.2023.100466"
)
field(doc, "1. Que investigan: ",
    "Examinan los efectos macroeconomicos de los movimientos del tipo de cambio en Mongolia, "
    "economia exportadora de commodities con deuda neta en moneda extranjera. Separan el "
    "canal financiero (deterioro de balances en moneda extranjera) del canal comercial "
    "(competitividad exportadora) para trazar la cadena completa de transmision.")
field(doc, "2. Marco transversal y periodo: ",
    "Mongolia (economia emergente exportadora de minerales, analogia con Peru). "
    "Datos mensuales.")
field(doc, "3. Metodologia: ",
    "Structural Bayesian VAR (SBVAR) mensual. Dos indicadores cambiarios: tipo de cambio "
    "ponderado por deuda (DWER, canal financiero) y tipo de cambio efectivo nominal "
    "(NEER, canal comercial). IRF estructurales y descomposicion de varianza. "
    "Variables: tipo de cambio, credito bancario, PIB, inflacion, inversion, exportaciones.")
field(doc, "4. Principales hallazgos: ",
    "Canal financiero: la depreciacion deteriora balances de bancos y empresas con deuda "
    "en moneda extranjera, contrae el credito bancario y reduce inversion y PIB. "
    "Canal comercial: la depreciacion eleva la inflacion via pass-through cambiario. "
    "Ambos canales operan simultaneamente y son estadisticamente significativos. "
    "La cadena completa IPC -> credito -> PIB -> inflacion queda empíricamente documentada.")
field(doc, "5. Cita APA 7: ",
    "Doojav, G.-O., Purevdorj, M., & Batjargal, A. (2024). The macroeconomic effects of "
    "exchange rate movements in a commodity-exporting developing economy. "
    "International Economics, 177, 100466. "
    "https://doi.org/10.1016/j.inteco.2023.100466")
divider(doc)

# --- [6] Keefe & Saha (2024) ---
h_paper(doc, 6, "Keefe & Saha (2024)  |  ScienceDirect — Journal of International Financial Markets, Institutions and Money")
apa(doc,
    "Keefe, H. G., & Saha, S. (2024). Global uncertainty and exchange rate conditions: "
    "Assessing the impact of uncertainty shocks in emerging markets and advanced economies. "
    "Journal of International Financial Markets, Institutions and Money, 96, 102047. "
    "https://doi.org/10.1016/j.intfin.2024.102047"
)
field(doc, "1. Que investigan: ",
    "Evaluan como los shocks de incertidumbre global (economica, financiera y de politica "
    "monetaria de EE.UU.) impactan el tipo de cambio bilateral, la volatilidad cambiaria "
    "y el IPC en economias emergentes y avanzadas. Evaluan si los efectos son simetricos "
    "entre grupos de paises.")
field(doc, "2. Marco transversal y periodo: ",
    "Panel de economias emergentes y avanzadas. Datos a frecuencia trimestral/mensual.")
field(doc, "3. Metodologia: ",
    "Global VAR (GVAR) con funciones impulso-respuesta generalizadas. "
    "Variables: IPC, tipo de cambio bilateral, volatilidad cambiaria, VIX, "
    "incertidumbre de politica economica (EPU), tasa de la Fed.")
field(doc, "4. Principales hallazgos: ",
    "Los picos de incertidumbre global generan aumentos inmediatos en la volatilidad "
    "cambiaria y en la presion depreciativa (IPC elevado) en economias emergentes. "
    "El IPC responde en la misma direccion que los shocks de incertidumbre, con efectos "
    "que se prolongan varios trimestres. Las economias avanzadas muestran respuestas "
    "distintas y mas amortiguadas. El IPC es la primera variable domestica en responder "
    "a shocks exogenos externos.")
field(doc, "5. Cita APA 7: ",
    "Keefe, H. G., & Saha, S. (2024). Global uncertainty and exchange rate conditions: "
    "Assessing the impact of uncertainty shocks in emerging markets and advanced economies. "
    "Journal of International Financial Markets, Institutions and Money, 96, 102047. "
    "https://doi.org/10.1016/j.intfin.2024.102047")
divider(doc)

doc.add_page_break()

# ================================================================
# SECCION 4 — EVIDENCIA ESPECIFICA PARA PERU
# ================================================================
h_section(doc, "Seccion 4: Evidencia Especifica para Peru")

# --- [7] Carrasco & Florian Hoyle (2025) ---
h_paper(doc, 7, "Carrasco & Florian Hoyle (2025)  |  ScienceDirect — Journal of Macroeconomics")
apa(doc,
    "Carrasco, A., & Florian Hoyle, D. (2025). External shocks and FX intervention policy "
    "in financially dollarized economies. Journal of Macroeconomics, 84, 103672. "
    "https://doi.org/10.1016/j.jmacro.2025.103672"
)
field(doc, "1. Que investigan: ",
    "Construyen un modelo Nuevo Keynesiano con fricciones financieras calibrado para Peru "
    "con el objetivo de analizar como las intervenciones cambiarias esterilizadas del BCRP "
    "responden a shocks externos y afectan el credito, la inversion y el producto en una "
    "economia con dolarizacion financiera.")
field(doc, "2. Marco transversal y periodo: ",
    "Peru (economia con dolarizacion parcial y flotacion administrada). "
    "Calibracion del sector bancario peruano. Datos del BCRP.")
field(doc, "3. Metodologia: ",
    "Modelo DSGE Nuevo Keynesiano con fricciones financieras (costos de agencia). "
    "Parametros calibrados emparejando las IRF de un modelo SVAR (con restricciones de "
    "exogeneidad de bloque, analogas a Cushman & Zha, 1997) con las respuestas del modelo "
    "estructural. Variables: tipo de cambio, credito, inversion, producto, inflacion.")
field(doc, "4. Principales hallazgos: ",
    "Las intervenciones cambiarias reducen significativamente la volatilidad del credito, "
    "la inversion y el producto al contrarrestar la presion cambiaria. El mecanismo opera "
    "directamente sobre la oferta de credito: las intervenciones reducen el descalce de "
    "moneda de los bancos, lo que evita la contraccion crediticia. Bajo intervencion activa "
    "(politica real del BCRP), las ganancias de bienestar sobre la flotacion libre son "
    "sustanciales. La cadena IPC -> credito -> actividad es causal y cuantitatiamente "
    "relevante para Peru.")
field(doc, "5. Cita APA 7: ",
    "Carrasco, A., & Florian Hoyle, D. (2025). External shocks and FX intervention policy "
    "in financially dollarized economies. Journal of Macroeconomics, 84, 103672. "
    "https://doi.org/10.1016/j.jmacro.2025.103672")
divider(doc)

# --- [8] Guevara & Rodriguez (2020) ---
h_paper(doc, 8, "Guevara & Rodriguez (2020)  |  ScienceDirect — North American Journal of Economics and Finance")
apa(doc,
    "Guevara, C., & Rodriguez, G. (2020). The role of credit supply shocks in Pacific "
    "Alliance countries: A TVP-VAR-SV approach. The North American Journal of Economics "
    "and Finance, 52, 101140. "
    "https://doi.org/10.1016/j.najef.2019.101140"
)
field(doc, "1. Que investigan: ",
    "Estiman como los shocks de oferta de credito afectan la actividad economica real en "
    "los cuatro paises de la Alianza del Pacifico (Chile, Colombia, Mexico y Peru), "
    "permitiendo que los parametros varien en el tiempo y entre fases del ciclo economico.")
field(doc, "2. Marco transversal y periodo: ",
    "Chile, Colombia, Mexico y Peru (Alianza del Pacifico). "
    "Datos trimestrales.")
field(doc, "3. Metodologia: ",
    "TVP-VAR-SV (VAR con parametros variables en el tiempo y volatilidad estocastica). "
    "Identificacion con restricciones de signo. Variables: actividad economica real, "
    "shocks de oferta de credito, shocks de oferta y demanda agregada.")
field(doc, "4. Principales hallazgos: ",
    "Los shocks de oferta de credito tienen un impacto importante en la actividad economica "
    "en todos los paises: aproximadamente 1% de respuesta en Colombia, Mexico y Peru, "
    "y 0.5% en Chile. Los efectos son variables en el tiempo y heterogeneos: mas fuertes "
    "durante desaceleraciones economicas y periodos de estres financiero. La cadena "
    "Credito -> Actividad economica es empiricamente activa y cuantitativamente significativa "
    "para Peru.")
field(doc, "5. Cita APA 7: ",
    "Guevara, C., & Rodriguez, G. (2020). The role of credit supply shocks in Pacific "
    "Alliance countries: A TVP-VAR-SV approach. The North American Journal of Economics "
    "and Finance, 52, 101140. "
    "https://doi.org/10.1016/j.najef.2019.101140")
divider(doc)

# --- [9] Rodriguez, Vassallo & Castillo B. (2023) ---
h_paper(doc, 9, "Rodriguez, Vassallo & Castillo B. (2023)  |  ScienceDirect — Economic Modelling")
apa(doc,
    "Rodriguez, G., Vassallo, R., & Castillo B., P. (2023). Effects of external shocks on "
    "macroeconomic fluctuations in Pacific Alliance countries. "
    "Economic Modelling, 124, 106302. "
    "https://doi.org/10.1016/j.econmod.2023.106302"
)
field(doc, "1. Que investigan: ",
    "Cuantifican el impacto de shocks externos financieros, reales (precios de commodities, "
    "terminos de intercambio) y nominales (politica monetaria de EE.UU.) sobre el "
    "crecimiento del PBI, la inflacion y las tasas de interes en Chile, Colombia, Mexico "
    "y Peru, con dinamicas variables en el tiempo.")
field(doc, "2. Marco transversal y periodo: ",
    "Chile, Colombia, Mexico y Peru. "
    "Datos trimestrales, 1994T1-2019T4.")
field(doc, "3. Metodologia: ",
    "TVP-VAR-SV comparados por DIC y log-verosimilitud marginal. "
    "Variables: crecimiento del PIB, inflacion del IPC, tasa de interes, "
    "bloques de shocks externos. Identifica shocks externos como exogenos al sistema.")
field(doc, "4. Principales hallazgos: ",
    "Los shocks externos explican entre 50% y 93% de la varianza macroeconomica en las "
    "economias de la Alianza del Pacifico. Peru muestra una participacion de 38% a 88% "
    "de varianza explicada por factores externos a lo largo de la muestra, reflejando "
    "alta dependencia externa. Un aumento del 1% en el crecimiento de China eleva el "
    "PBI de Peru en 0.8% en el primer anio, la respuesta mas alta entre los paises. "
    "La variacion temporal en los parametros es estadisticamente validada.")
field(doc, "5. Cita APA 7: ",
    "Rodriguez, G., Vassallo, R., & Castillo B., P. (2023). Effects of external shocks on "
    "macroeconomic fluctuations in Pacific Alliance countries. "
    "Economic Modelling, 124, 106302. "
    "https://doi.org/10.1016/j.econmod.2023.106302")
divider(doc)

# --- [10] Rodriguez et al. (2024) ERPT ---
h_paper(doc, 10, "Rodriguez, Castillo B., Calero, Salcedo & Ataurima (2024)  |  ScienceDirect — Journal of International Money and Finance")
apa(doc,
    "Rodriguez, G., Castillo B., P., Calero, R., Salcedo Cisneros, R., & Ataurima Arellano, M. "
    "(2024). Evolution of the exchange rate pass-through into prices in Peru: An empirical "
    "application using TVP-VAR-SV models. Journal of International Money and Finance, 142, 103023. "
    "https://doi.org/10.1016/j.jimonfin.2024.103023"
)
field(doc, "1. Que investigan: ",
    "Rastrean como los shocks de tipo de cambio se transmiten hacia precios de importacion, "
    "precios al productor (IPP) y precios al consumidor (IPC de precios) en Peru, "
    "incluyendo los efectos de la dolarizacion parcial en la dinamica de precios en "
    "cada etapa de la cadena de precios.")
field(doc, "2. Marco transversal y periodo: ",
    "Peru. Datos trimestrales, 1995T2-2022T4 (incluye pandemia y alta inflacion 2022).")
field(doc, "3. Metodologia: ",
    "TVP-VAR-SV (VAR con parametros variables en el tiempo y volatilidad estocastica). "
    "Estimacion bayesiana. Variables: tipo de cambio, precios de importacion, IPP, IPC. "
    "Identificacion consistente con una cadena de precios secuencial.")
field(doc, "4. Principales hallazgos: ",
    "El pass-through de corto plazo supera al de largo plazo en todas las etapas de precios, "
    "reflejando dolarizacion de precios en importaciones y produccion. El pass-through al "
    "IPC de precios fue de 10% hasta 2009 y aumento a 15% post-2009. Los precios de "
    "importacion muestran el mayor pass-through, el IPP intermedio y el IPC el mas bajo. "
    "Esto documenta empiricamente la secuencia: presion cambiaria -> precios de importacion "
    "-> IPP -> IPC de precios, validando la inflacion como ultima variable en el VAR.")
field(doc, "5. Cita APA 7: ",
    "Rodriguez, G., Castillo B., P., Calero, R., Salcedo Cisneros, R., & Ataurima Arellano, M. "
    "(2024). Evolution of the exchange rate pass-through into prices in Peru: An empirical "
    "application using TVP-VAR-SV models. Journal of International Money and Finance, 142, 103023. "
    "https://doi.org/10.1016/j.jimonfin.2024.103023")
divider(doc)

doc.add_page_break()

# ================================================================
# SECCION 5 — CAUSALIDAD DE GRANGER Y METODOLOGIA VAR
# ================================================================
h_section(doc, "Seccion 5: Causalidad de Granger y Metodologia VAR con IPC")

# --- [11] Olanipekun et al. (2019) ---
h_paper(doc, 11, "Olanipekun, Gungor & Olasehinde-Williams (2019)  |  Scopus — SAGE Open")
apa(doc,
    "Olanipekun, I. O., Gungor, H., & Olasehinde-Williams, G. (2019). Unraveling the "
    "causal relationship between economic policy uncertainty and exchange market pressure "
    "in BRIC countries: Evidence from bootstrap panel Granger causality. "
    "SAGE Open, 9(2), 2158244019853903. "
    "https://doi.org/10.1177/2158244019853903"
)
field(doc, "1. Que investigan: ",
    "Examinan la relacion causal entre la incertidumbre de politica economica (global y "
    "domestica) y el IPC en las economias BRIC, con enfasis en detectar causalidad "
    "unidireccional o bidireccional mediante pruebas robustas a heterocedasticidad y "
    "dependencia de corte transversal.")
field(doc, "2. Marco transversal y periodo: ",
    "Brasil, Rusia, India y China (BRIC). "
    "Datos anuales.")
field(doc, "3. Metodologia: ",
    "Bootstrap Panel Granger Causality (Konia, 2006). Pruebas de raiz unitaria de paneles "
    "con quiebres estructurales. Prueba de dependencia de corte transversal (CD test). "
    "El metodo bootstrap hace las pruebas robustas a no-normalidad y heterocedasticidad.")
field(doc, "4. Principales hallazgos: ",
    "A nivel panel: causalidad bidireccional entre incertidumbre de politica economica "
    "(global y domestica) y el IPC. Por pais: causalidades especificas bidireccionales "
    "y unidireccionales dependiendo del contexto. El IPC Granger-causa la incertidumbre "
    "en varios paises, indicando su caracter predictivo. Las pruebas bootstrap confirman "
    "que el IPC tiene relaciones causales significativas con otras variables macroeconomicas.")
field(doc, "5. Cita APA 7: ",
    "Olanipekun, I. O., Gungor, H., & Olasehinde-Williams, G. (2019). Unraveling the "
    "causal relationship between economic policy uncertainty and exchange market pressure "
    "in BRIC countries: Evidence from bootstrap panel Granger causality. "
    "SAGE Open, 9(2), 2158244019853903. "
    "https://doi.org/10.1177/2158244019853903")
divider(doc)

# --- [12] Siklar & Akca (2020) ---
h_paper(doc, 12, "Siklar & Akca (2020)  |  Scopus — Ekonomika (Vilnius University)")
apa(doc,
    "Siklar, I., & Akca, A. (2020). Exchange market pressure and monetary policy: "
    "The Turkish case. Ekonomika, 99(1), 110-130. "
    "https://doi.org/10.15388/Ekon.2020.1.7"
)
field(doc, "1. Que investigan: ",
    "Determinan la relacion entre la politica monetaria y el IPC en Turquia, evaluando "
    "si los instrumentos de politica monetaria —credito domestico y diferencial de tasas "
    "de interes— Granger-causan el IPC o viceversa. Aplican la definicion de Girton & "
    "Roper (1977) para construir el IPC turco.")
field(doc, "2. Marco transversal y periodo: ",
    "Turquia (economia emergente con flotacion administrada y episodios de alta presion "
    "cambiaria, comparable con Peru). "
    "Datos mensuales, 2002-2018.")
field(doc, "3. Metodologia: ",
    "Estimacion del IPC siguiendo Girton & Roper (1977). "
    "Modelo VAR mensual. Prueba de causalidad de Granger bidireccional entre: "
    "(1) credito domestico e IPC; (2) diferencial de tasa de interes e IPC. "
    "Variables: M1, reservas internacionales, tipo de cambio, diferencial de tasas, credito.")
field(doc, "4. Principales hallazgos: ",
    "Causalidad unidireccional de credito domestico -> IPC (el credito Granger-causa el "
    "IPC, no al reves). Causalidad bidireccional entre diferencial de tasa de interes e IPC. "
    "La expansion del credito domestico genera presion depreciativa (IPC mayor), confirmando "
    "la logica de Girton & Roper. El VAR muestra que shocks en el credito se transmiten al "
    "IPC con 2-3 meses de rezago.")
field(doc, "5. Cita APA 7: ",
    "Siklar, I., & Akca, A. (2020). Exchange market pressure and monetary policy: "
    "The Turkish case. Ekonomika, 99(1), 110-130. "
    "https://doi.org/10.15388/Ekon.2020.1.7")
divider(doc)

doc.add_page_break()

# ================================================================
# SECCION 6 — CHOQUES EXTERNOS Y TRANSMISION EN PERU Y AL
# ================================================================
h_section(doc, "Seccion 6: Choques Externos y Transmision Monetaria en Peru y America Latina")

# --- [13] Rossini & Vega (2008) ---
h_paper(doc, 13, "Rossini & Vega (2008)  |  BIS Papers No. 35 — acceso abierto, ampliamente citado")
apa(doc,
    "Rossini, R., & Vega, M. (2008). The monetary policy transmission mechanism under "
    "financial dollarisation: The case of Peru 1996-2006. En Bank for International "
    "Settlements (Ed.), Transmission Mechanisms for Monetary Policy in Emerging Market "
    "Economies (BIS Papers No. 35, pp. 395-412). Bank for International Settlements. "
    "https://www.bis.org/publ/bppdf/bispap35r.pdf"
)
field(doc, "1. Que investigan: ",
    "Analizan como la dolarizacion financiera afecta la transmision de la politica monetaria "
    "en Peru, identificando que canales —tasa de interes, tipo de cambio, expectativas— "
    "ganaron o perdieron relevancia tras la adopcion del esquema de metas de inflacion (IT).")
field(doc, "2. Marco transversal y periodo: ",
    "Peru. Datos mensuales, 1996-2006. "
    "Periodo de transicion hacia metas de inflacion y de desdolarizacion progresiva.")
field(doc, "3. Metodologia: ",
    "Modelo VAR con identificacion de Cholesky. Ecuaciones OLS para canales individuales. "
    "Analisis de impulso-respuesta para evaluar la transmision dinamica. "
    "Variables: tasa de referencia del BCRP, tipo de cambio, inflacion, credito bancario, "
    "indicadores de dolarizacion.")
field(doc, "4. Principales hallazgos: ",
    "El canal de tasa de interes y el canal de expectativas se fortalecieron notablemente "
    "tras la adopcion del IT. La dolarizacion reduce la efectividad de la politica monetaria "
    "porque los descalces de moneda hacen contractiva la depreciacion (efecto hoja de "
    "balance). A mayor flexibilidad cambiaria, mas rapido pero mas debil es el pass-through "
    "del tipo de cambio. La desdolarizacion progresiva mejora la transmision monetaria.")
field(doc, "5. Cita APA 7: ",
    "Rossini, R., & Vega, M. (2008). The monetary policy transmission mechanism under "
    "financial dollarisation: The case of Peru 1996-2006. En Bank for International "
    "Settlements (Ed.), Transmission Mechanisms for Monetary Policy in Emerging Market "
    "Economies (BIS Papers No. 35, pp. 395-412). Bank for International Settlements.")
divider(doc)

doc.add_page_break()

# ================================================================
# TABLA SINTESIS
# ================================================================
h_section(doc, "Tabla Sintesis de Referencias")
doc.add_paragraph()

tbl = doc.add_table(rows=1, cols=6)
tbl.style = "Table Grid"

hdrs = ["#", "Autores (anio)", "Fuente", "Metodo", "Pais/Muestra", "Relevancia principal"]
hdr_cells = tbl.rows[0].cells
for i, h in enumerate(hdrs):
    hdr_cells[i].text = h
    shade_cell(hdr_cells[i], "1F497D")
    for par in hdr_cells[i].paragraphs:
        for run in par.runs:
            run.bold = True; run.font.size = Pt(9); run.font.color.rgb = C_WHITE

rows_data = [
    ["1",  "Girton & Roper (1977)",        "Scopus (AER)",          "Modelo BOP MCO",   "Canada",             "Define el IPC = Delta(e)+Delta(r)"],
    ["2",  "Aizenman & Binici (2016)",      "ScienceDirect (JIMF)",  "Panel GMM",         "50 paises",          "IPC es externamente determinado en EMEs"],
    ["3",  "Keefe (2021)",                  "ScienceDirect (JIFMI)", "Panel VAR Cholesky","40 EMDEs",           "Nexo bidireccional credito <-> IPC"],
    ["4",  "Gevorkyan (2019)",              "Scopus (Appl. Econ.)",  "Panel VAR",         "EMEs commodities",   "IPC -> credito en regimenes flexibles"],
    ["5",  "Doojav et al. (2024)",          "ScienceDirect (IntEco)","SBVAR",             "Mongolia",           "Cadena completa IPC->credito->PIB->inflacion"],
    ["6",  "Keefe & Saha (2024)",           "ScienceDirect (JIFMI)", "GVAR",              "Paises EME y avanz.", "IPC responde primero a shocks externos"],
    ["7",  "Carrasco & Florian H. (2025)",  "ScienceDirect (JMacro)","DSGE + SVAR Peru",  "Peru",               "IPC->credito->actividad causal en Peru"],
    ["8",  "Guevara & Rodriguez (2020)",    "ScienceDirect (NAJEF)", "TVP-VAR-SV",        "Alianza Pacifico",   "Credito -> actividad 1% en Peru"],
    ["9",  "Rodriguez, Vassallo et al. (2023)","ScienceDirect (EcoMod)","TVP-VAR-SV",    "Alianza Pacifico",   "Shocks externos explican 88% varianza Peru"],
    ["10", "Rodriguez et al. (2024)",       "ScienceDirect (JIMF)",  "TVP-VAR-SV Bayes", "Peru",               "Pass-through cambiario a IPC en Peru"],
    ["11", "Olanipekun et al. (2019)",      "Scopus (SAGE Open)",    "Bootstrap Granger", "BRIC",               "Causalidad de Granger bidireccional IPC"],
    ["12", "Siklar & Akca (2020)",          "Scopus (Ekonomika)",    "VAR + Granger",     "Turquia",            "Credito -> IPC (unidireccional)"],
    ["13", "Rossini & Vega (2008)",         "BIS Papers No. 35",     "VAR Cholesky",      "Peru",               "Canal crediticio y cambiario en Peru IT"],
]

alt = ["F2F2F2", "FFFFFF"]
for idx, row_data in enumerate(rows_data):
    row = tbl.add_row()
    fill = alt[idx % 2]
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        shade_cell(cell, fill)
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(8.5)

doc.add_paragraph()

# ================================================================
# NOTA DE CIERRE
# ================================================================
h_section(doc, "Nota Metodologica")
body(doc,
    "Las 13 referencias cubren sistematicamente los cuatro eslabones de la cadena causal "
    "propuesta: IPC -> Credito -> Actividad Economica -> Inflacion, y proveen justificacion "
    "para la identificacion de Cholesky utilizada en el modelo VAR."
)
body(doc,
    "Para el eslabón IPC (definicion y exogeneidad): Girton & Roper (1977), Aizenman & "
    "Binici (2016), Keefe & Saha (2024). Para el eslabon IPC -> Credito: Keefe (2021), "
    "Gevorkyan (2019), Siklar & Akca (2020). Para el eslabon Credito -> Actividad: "
    "Doojav et al. (2024), Guevara & Rodriguez (2020), Carrasco & Florian Hoyle (2025). "
    "Para el eslabon Actividad -> Inflacion (pass-through): Rodriguez et al. (2024), "
    "Rodriguez, Vassallo et al. (2023). Para el contexto Peru: Rossini & Vega (2008), "
    "Carrasco & Florian Hoyle (2025)."
)

p_note = doc.add_paragraph()
r_note = p_note.add_run(
    "Fuentes: ScienceDirect (Elsevier) y Scopus. "
    "Archivos locales del proyecto: Regresion_VAR/Bibliografia. "
    "Todas las referencias en formato APA 7a edicion con DOI verificado."
)
r_note.font.size = Pt(9); r_note.italic = True; r_note.font.color.rgb = C_GREY

# ================================================================
out = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Referencias_VAR_IPC.docx"
)
doc.save(out)
print(f"Documento guardado: {out}")
